from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from enum import StrEnum

from docx import Document

from .entities import FinancialEntity, CANONICAL_ORDER


@dataclass(frozen=True)
class ExtractionResult:
    entities: dict[str, str]  # Only the 9 target entities
    evidence: dict[str, list[str]]

    def to_json_dict(self) -> dict[str, Any]:
        return {"entities": self.entities, "evidence": self.evidence}


# Patterns for value type detection
DATE_PATTERNS = [
    r"\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}",
    r"\d{1,2}/\d{1,2}/\d{2,4}",
    r"\d{4}-\d{2}-\d{2}",
]

CURRENCY_PATTERNS = [
    r"(?:EUR|USD|GBP|JPY|CHF)\s+\d+(?:\.\d+)?\s*(?:million|billion|thousand|M|B|K)?",
    r"[€$£]\s?\d+(?:[.,]\d+)?\s*(?:million|billion|M|B)?",
]

PERCENTAGE_PATTERNS = [
    r"\d+(?:\.\d+)?\s*%",
]

ISIN_PATTERN = r"[A-Z]{2}[A-Z0-9]{9}[0-9]"


def _looks_like_date(value: str) -> bool:
    return any(re.search(p, value, re.IGNORECASE) for p in DATE_PATTERNS)


def _looks_like_currency(value: str) -> bool:
    return any(re.search(p, value, re.IGNORECASE) for p in CURRENCY_PATTERNS)


def _looks_like_percentage(value: str) -> bool:
    return any(re.search(p, value, re.IGNORECASE) for p in PERCENTAGE_PATTERNS)


def _looks_like_isin(value: str) -> bool:
    return bool(re.search(ISIN_PATTERN, value))


def _normalize_key(key: str) -> str:
    """Clean up key for better matching."""
    k = key.strip().lower()
    # Remove parentheses and their content
    k = re.sub(r"\s*\([^)]*\)", "", k)
    # Replace multiple spaces with single
    k = re.sub(r"\s+", " ", k)
    k = k.replace("\u00a0", " ")
    return k.strip()


def _infer_entity_type(key: str, value: str) -> FinancialEntity | None:
    """Map a key-value pair to one of the 9 FinancialEntity types."""
    key_lower = _normalize_key(key)
    value_lower = value.lower()
    
    # Rule-based mapping to target entities
    # Counterparty
    if any(term in key_lower for term in ["party", "counterparty", "cp"]):
        return FinancialEntity.counterparty
    
    # Initial Valuation Date
    if any(term in key_lower for term in ["initial valuation", "initial val", "pricing date"]):
        return FinancialEntity.initial_valuation_date
    
    # Valuation Date
    if any(term in key_lower for term in ["valuation date", "val date"]) and "initial" not in key_lower:
        return FinancialEntity.valuation_date
    
    # Maturity
    if any(term in key_lower for term in ["maturity", "termination date", "end date"]):
        return FinancialEntity.maturity
    
    # Notional
    if any(term in key_lower for term in ["notional", "nominal", "face value"]):
        return FinancialEntity.notional
    
    # Underlying
    if any(term in key_lower for term in ["underlying", "reference asset", "linked to"]):
        return FinancialEntity.underlying
    
    # Coupon
    if any(term in key_lower for term in ["coupon", "interest rate", "rate"]):
        return FinancialEntity.coupon
    
    # Barrier
    if any(term in key_lower for term in ["barrier", "knock", "trigger"]):
        return FinancialEntity.barrier
    
    # Calendar
    if any(term in key_lower for term in ["calendar", "business day", "day count"]):
        return FinancialEntity.calendar
    
    # Fallback: try to infer from value type
    if _looks_like_date(value):
        # Could be initial valuation, valuation, or maturity
        if "initial" in key_lower or "pricing" in key_lower:
            return FinancialEntity.initial_valuation_date
        elif "valuation" in key_lower:
            return FinancialEntity.valuation_date
        elif "maturity" in key_lower or "termination" in key_lower:
            return FinancialEntity.maturity
        elif "trade" in key_lower or "effective" in key_lower:
            # Not in our target list, but still useful — skip
            return None
        else:
            # Generic date — don't know which entity
            return None
    
    if _looks_like_currency(value):
        return FinancialEntity.notional
    
    if _looks_like_percentage(value):
        if "coupon" in key_lower or "rate" in key_lower:
            return FinancialEntity.coupon
        elif "barrier" in key_lower:
            return FinancialEntity.barrier
        return FinancialEntity.coupon  # Default to coupon
    
    if _looks_like_isin(value):
        # ISIN is part of underlying
        return FinancialEntity.underlying
    
    return None


def _extract_kv_from_line(line: str) -> tuple[str, str] | None:
    """Extract key-value pair using various separators."""
    separators = [r"\s*:\s*", r"\s*=\s*", r"\s*-\s*", r"\s*–\s*"]
    
    for sep in separators:
        parts = re.split(sep, line, maxsplit=1)
        if len(parts) == 2:
            key, value = parts[0].strip(), parts[1].strip()
            if key and value and len(key) < 50:
                return key, value
    
    return None


def _iter_doc_content(doc: Document) -> list[str]:
    """Extract all text lines from paragraphs and tables."""
    lines = []
    
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            lines.append(text)
    
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if not cells:
                continue
            if len(cells) == 1:
                lines.append(cells[0])
            elif len(cells) == 2:
                lines.append(f"{cells[0]}: {cells[1]}")
            else:
                lines.append(" | ".join(cells))
    
    return lines


def _extract_from_doc(doc: Document) -> ExtractionResult:
    lines = _iter_doc_content(doc)
    
    # Initialize with None for all target entities
    entities: dict[str, str] = {entity.value: "" for entity in CANONICAL_ORDER}
    evidence: dict[str, list[str]] = {entity.value: [] for entity in CANONICAL_ORDER}
    
    for line in lines:
        kv = _extract_kv_from_line(line)
        if not kv:
            continue
        
        raw_key, raw_value = kv
        entity_type = _infer_entity_type(raw_key, raw_value)
        
        if entity_type and not entities[entity_type.value]:  # Only take first occurrence
            entities[entity_type.value] = raw_value
            evidence[entity_type.value].append(line)
    
    # Remove empty evidence
    evidence = {k: v for k, v in evidence.items() if v}
    
    return ExtractionResult(entities=entities, evidence=evidence)


def extract_entities_from_docx_bytes(data: bytes) -> ExtractionResult:
    import io
    doc = Document(io.BytesIO(data))
    return _extract_from_doc(doc)